import re
import docx

class ASTNode:
    pass

class TextNode(ASTNode):
    
    def __init__(self, text: str):
        
        self.text = text

class CompositeNode(ASTNode):
    
    def __init__(self, children=None):
        
        self.children = children if children is not None else []

class BoldNode(CompositeNode): pass
class ItalicNode(CompositeNode): pass

class MarkdownDocxBuilder:
    
    def __init__(self):
        # Lexer Rule. 
        # Searches for ** at the beginning and ** at the end if ** exists, 
        # using the "\1" reference. If * exists, it searches for *.
        self.md_pattern = re.compile(r'(\*\*|\*)(.*?)\1')

    def _parse_text(self, text: str) -> list:
        
        if not text:
            return []

        match = self.md_pattern.search(text)
        
        # If no Markdown tags remain, text is a pure TextNode.
        if not match:
            return [TextNode(text)]

        nodes = []
        start_idx = match.start()
        end_idx = match.end()
        tag = match.group(1)      # '**' or '*'
        content = match.group(2)  # Text within the tag

        # Prefix: Plain text before the match.
        if start_idx > 0:
            nodes.append(TextNode(text[:start_idx]))

        # Content: Inside the match. Recursion for nested structures.
        inner_nodes = self._parse_text(content)
        if tag == '**':
            nodes.append(BoldNode(inner_nodes))
        elif tag == '*':
            nodes.append(ItalicNode(inner_nodes))

        # 3. Suffix: Text remaining after matching. Recursion continues for the remaining text.
        if end_idx < len(text):
            nodes.extend(self._parse_text(text[end_idx:]))

        return nodes

    def _render_nodes(self, nodes: list, paragraph, current_state: dict = None):
        """
        Traversal AST Tree from top to bottom. Creates run objects by transferring states.
        """
        if current_state is None:
            current_state = {'bold': False, 'italic': False}

        for node in nodes:
            # Leaf Node. Apply State and print Run object to Word
            if isinstance(node, TextNode):
                if node.text: # Check to avoid creating empty runways
                    run = paragraph.add_run(node.text)
                    run.bold = current_state.get('bold', False)
                    run.italic = current_state.get('italic', False)
            
            # Branch Knot. Add format to the state.
            elif isinstance(node, CompositeNode):
                new_state = current_state.copy()
                
                if isinstance(node, BoldNode):
                    new_state['bold'] = True
                elif isinstance(node, ItalicNode):
                    new_state['italic'] = True
                    
                self._render_nodes(node.children, paragraph, new_state)

    def add_markdown_paragraph(self, doc, text: str):
        
        paragraph = doc.add_paragraph()
        ast_nodes = self._parse_text(text)
        self._render_nodes(ast_nodes, paragraph)
        
        return paragraph

if __name__ == "__main__":
    
    doc = docx.Document()
    builder = MarkdownDocxBuilder()
    
    test_metni = "Sistem **hızlı ve *güvenli* hale** geldi. İşlem tamam."
    
    builder.add_markdown_paragraph(doc, test_metni)
    doc.save("test_output.docx")