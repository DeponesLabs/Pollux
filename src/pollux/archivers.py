import os
import re
from datetime import datetime
import json
from abc import ABC, abstractmethod
from docx import Document
from docx.shared import RGBColor, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

class BaseArchiver(ABC):
        
    def __init__(self, filename):
    
        self.filename = filename        

    @abstractmethod
    def append_message(self, role: str, text: str):
        pass

    @abstractmethod
    def save(self):
        pass

class DocxArchiver(BaseArchiver):
    
    def __init__(self, filename, heading='Pollux AI - Session Record', heading_alignment: WD_ALIGN_PARAGRAPH = WD_ALIGN_PARAGRAPH.CENTER):
        
        super().__init__(f"{filename}.docx")
        
        self.doc = Document()
        self.add_heading(heading, 0, heading_alignment)
        self.add_document_date()
        self.doc.add_paragraph()

    def add_heading(self, heading, level, alignment: WD_ALIGN_PARAGRAPH = WD_ALIGN_PARAGRAPH.CENTER):
        
        heading = self.doc.add_heading(heading, level)
        heading.alignment = alignment
    
    def add_document_date(self):
        
        timestamp = datetime.now().strftime('%Y-%m-%d')
        session_date = f"Session Date: {timestamp}"
        self.add_paragraph(text=session_date, alignment=WD_ALIGN_PARAGRAPH.RIGHT, bold=True, underline=True)
    
    def add_paragraph(self, text: str, alignment: WD_ALIGN_PARAGRAPH = WD_ALIGN_PARAGRAPH.JUSTIFY,
                      bold: bool = False, italic: bool = False, underline = False,  
                      strike = False, subscript = False, superscript = False):
        
        date_p = self.doc.add_paragraph()
        date_p.alignment = alignment
        run = date_p.add_run(text)

        run.font.bold = bold
        run.font.italic = italic
        run.font.underline = underline
        run.font.strike = strike
        run.font.subscript = subscript
        run.font.superscript = superscript

    def append_message(self, role: str, text: str):
        
        timestamp = datetime.now().strftime("%H:%M:%S")
        p = self.doc.add_paragraph()
        
        label_run = p.add_run(f"[{timestamp}] {role}:\n")
        label_run.bold = True
        
        text_run = p.add_run(text)
        
        if role == "User":
            # Messages of User
            label_run.font.color.rgb = RGBColor(0, 51, 153)
            text_run.font.color.rgb = RGBColor(0, 51, 153)
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT    # Align right as Chat Apps do
        else:
            # Messages of Pollux
            label_run.font.color.rgb = RGBColor(34, 139, 34) # Forest Green
            text_run.font.size = Pt(11)
            
        self.doc.add_paragraph("_" * 60)
        
        self.save()

    def save(self):
        
        self.doc.save(self.filename)

class MdArchiver(BaseArchiver):
    
    def __init__(self, filename):
        
        super().__init__(f"{filename}.md")

    def append_message(self, role: str, text: str):
        
        timestamp = datetime.now().strftime("%H:%M:%S")
            
        with open(self.filename, "a", encoding="utf-8") as f:
            f.write(f"### [{timestamp}] {role}\n{text}\n\n---\n\n")

    def append_block(self, text: str):
        
        with open(self.filename, "a", encoding="utf-8") as f:
            f.write(f"### {text}\n\n---\n\n")

    def save(self):
        pass

class TxtArchiver(BaseArchiver):
    
    def __init__(self, filename):
        
        super().__init__(f"{filename}.txt")
        
    def _strip_markdown(self, raw_text: str) -> str:

        # Cleaning order is important!
        # Links: Extract the anchor text, drop the URL -> [text](url) to text
        clean_text = re.sub(r'\[([^\]]+)\]\([^\)]+\)', r'\1', raw_text)

        # Code blocks & Inline code: Remove backticks and language identifiers
        clean_text = re.sub(r'```[a-zA-Z]*\n?|`', '', clean_text)

        # Bold & Italic: Capture the text inside, drop the asterisks/underscores
        clean_text = re.sub(r'[\*\_]{1,3}([^\*\_]+)[\*\_]{1,3}', r'\1', clean_text)

        # Line Prefixes: Remove Headers, Blockquotes, and List markers (using MULTILINE flag)
        clean_text = re.sub(r'^(#+|>|[\*\-\+]|\d+\.)\s+', '', clean_text, flags=re.MULTILINE)

        return clean_text
    
    def append_message(self, role, text):
        
        timestamp = datetime.now().strftime("%H:%M:%S")
        
        # Clean the text BEFORE writing to I/O
        clean_text = self._strip_markdown(text)
        
        with open(self.filename, "a", encoding="utf-8") as f:
            f.write(f"[{timestamp}] {role.upper()}: {clean_text}\n")
            f.write("-" * 40 + "\n")
    

class JsonArchiver(BaseArchiver):
    
    def __init__(self, filename):
        
        super().__init__(f"{filename}.json")
        
        self.history = []

    def append_message(self, role: str, text: str):
        
        self.history.append({
            "timestamp": datetime.now().isoformat(),
            "role": role,
            "text": text
        })

    def save(self):
        
        with open(self.filename, "w", encoding="utf-8") as f:
            json.dump(self.history, f, indent=4, ensure_ascii=False)
