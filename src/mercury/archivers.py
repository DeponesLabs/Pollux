import os
from datetime import datetime
import json
from abc import ABC, abstractmethod
from docx import Document
from docx.shared import RGBColor, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

class BaseArchiver(ABC):
    
    def __init__(self, base_filename):
    
        self.filename = base_filename

    @abstractmethod
    def append_message(self, role: str, text: str):
        pass

    @abstractmethod
    def save(self):
        pass


class DocxArchiver(BaseArchiver):
    
    def __init__(self, base_filename, heading='Pollux AI - Session Record', level=0):
        
        super().__init__(f"{base_filename}.docx")
        
        self.doc = Document()
        self.doc.add_heading(heading, level)

    def append_message(self, role: str, text: str):
        
        timestamp = datetime.now().strftime("%H:%M:%S")
        
        p = self.doc.add_paragraph()
        
        p.add_run(f"[{timestamp}] {role}: ").bold = True
        p.add_run(text)
        
        self.save() 

    def save(self):
        
        self.doc.save(self.filename)

class DocxArchiver2(BaseArchiver):
    def __init__(self, base_filename):
        super().__init__(f"{base_filename}.docx")
        self.doc = Document()
        
        # 1. Ana Başlık (Ortalanmış ve Şık)
        heading = self.doc.add_heading('🌟 Pollux AI - Session Record', 0)
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # 2. Alt Bilgi (Tarih)
        date_p = self.doc.add_paragraph(f"Session Date: {datetime.now().strftime('%Y-%m-%d')}")
        date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        date_p.runs[0].font.italic = True
        
        self.doc.add_paragraph() # Boşluk bırak

    def append_message(self, role: str, text: str):
        timestamp = datetime.now().strftime("%H:%M:%S")
        p = self.doc.add_paragraph()
        
        # 1. Etiket Kısmı (Saat ve Kimin Yazdığı)
        label_run = p.add_run(f"[{timestamp}] {role}:\n")
        label_run.bold = True
        
        # 2. Metin Kısmı
        text_run = p.add_run(text)
        
        # 3. Rol tabanlı renklendirme ve stil (İşte büyü burada!)
        if role == "User":
            # Kullanıcı mesajları: Koyu Mavi
            label_run.font.color.rgb = RGBColor(0, 51, 153)
            text_run.font.color.rgb = RGBColor(0, 51, 153)
            # p.alignment = WD_ALIGN_PARAGRAPH.RIGHT # İstersen WhatsApp gibi sağa yaslayabilirsin!
        else:
            # Pollux mesajları: Etiket Yeşil, Metin Siyah
            label_run.font.color.rgb = RGBColor(34, 139, 34) # Orman Yeşili
            # Yapay zeka metni daha okunaklı olsun diye fontu bir tık büyütebiliriz
            text_run.font.size = Pt(11)
            
        # Araya ayırıcı bir çizgi (Word stili)
        self.doc.add_paragraph("_" * 60)
        
        # Her adımda güvenle kaydet
        self.save()

    def save(self):
        self.doc.save(self.filename)

class TxtArchiver(BaseArchiver):
    
    def __init__(self, base_filename):
        
        super().__init__(f"{base_filename}.txt")

    def append_message(self, role: str, text: str):
        
        timestamp = datetime.now().strftime("%H:%M:%S")
        
        with open(self.filename, "a", encoding="utf-8") as f:
            f.write(f"[{timestamp}] {role}: {text}\n")
            f.write("-" * 40 + "\n")

    def save(self):
        pass


class JsonArchiver(BaseArchiver):
    
    def __init__(self, base_filename):
        
        super().__init__(f"{base_filename}.json")
        
        self.history = []

    def append_message(self, role: str, text: str):
        
        self.history.append({
            "timestamp": datetime.now().isoformat(),
            "role": role,
            "text": text
        })

    def save(self):
        
        with open(self.filename, "w", encoding="utf-8") as f:
            json.dump(self.history, f, indent=4, ensure_ascii=False)
